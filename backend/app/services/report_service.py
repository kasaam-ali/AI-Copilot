"""Generate downloadable PDF / DOCX inspection reports.

Charts are rendered server-side with matplotlib (Agg); the PDF is built with ReportLab and
the DOCX with python-docx. Both consume the same gathered evidence so they stay in sync.
"""

from __future__ import annotations

import io
import json
from datetime import datetime, timezone
from pathlib import Path

from sqlmodel import Session, select

from app.config import get_settings
from app.models_db.models import HitlDecision, Inspection, Prediction
from app.services.analysis_service import analyze_inspection

_BAND_COLOR = {
    "healthy": "#10b981",
    "watch": "#f5a623",
    "at_risk": "#f97316",
    "defect": "#ef4444",
    "unknown": "#6b7280",
}


def _hbar_png(labels: list[str], values: list[float], colors: list[str], title: str) -> bytes:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(5.2, max(1.6, 0.42 * len(labels) + 0.8)))
    y = range(len(labels))
    ax.barh(list(y), values, color=colors)
    ax.set_yticks(list(y))
    ax.set_yticklabels([label.replace("_", " ") for label in labels], fontsize=9)
    ax.invert_yaxis()
    ax.axvline(0, color="#9ca3af", linewidth=0.8)
    ax.set_title(title, fontsize=11, loc="left")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    buffer = io.BytesIO()
    fig.savefig(buffer, format="png", dpi=140)
    plt.close(fig)
    return buffer.getvalue()


def _health_png(score: float | None, band: str) -> bytes:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(5.2, 1.2))
    color = _BAND_COLOR.get(band, _BAND_COLOR["unknown"])
    ax.barh([0], [100], color="#e5e7eb", height=0.5)
    if score is not None:
        ax.barh([0], [score], color=color, height=0.5)
        ax.text(min(score, 92), 0, f" {score:.0f}", va="center", fontsize=13, fontweight="bold")
    ax.set_xlim(0, 100)
    ax.set_yticks([])
    ax.set_title(f"Fused health score — {band.replace('_', ' ')}", fontsize=11, loc="left")
    for spine in ax.spines.values():
        spine.set_visible(False)
    fig.tight_layout()
    buffer = io.BytesIO()
    fig.savefig(buffer, format="png", dpi=140)
    plt.close(fig)
    return buffer.getvalue()


def _narrative(session: Session, inspection: Inspection) -> dict:
    if not inspection.llm_narrative:
        analyze_inspection(session, inspection.id)
        session.refresh(inspection)
    try:
        return json.loads(inspection.llm_narrative) if inspection.llm_narrative else {}
    except json.JSONDecodeError:
        return {}


def _gather(session: Session, inspection_id: int) -> dict | None:
    inspection = session.get(Inspection, inspection_id)
    if inspection is None:
        return None
    predictions = session.exec(
        select(Prediction).where(Prediction.inspection_id == inspection_id)
    ).all()
    decisions = session.exec(
        select(HitlDecision).where(HitlDecision.inspection_id == inspection_id)
    ).all()

    charts: dict[str, bytes] = {}
    gradcam_path: str | None = None
    stats: list[tuple[str, str]] = []

    for pred in predictions:
        kind = pred.model_type.value
        if kind == "image":
            gradcam_path = pred.artifact_path
            stats.append(("Vision", f"{pred.output.get('label')} · {pred.raw_scores.get('defect_probability', 0):.0%} defect prob"))
        elif kind == "tabular":
            shap = pred.raw_scores.get("shap", [])[:8]
            if shap:
                charts["tabular"] = _hbar_png(
                    [d["feature"] for d in shap],
                    [d["contribution"] for d in shap],
                    ["#ef4444" if d["contribution"] >= 0 else "#10b981" for d in shap],
                    "Process-data SHAP contributions",
                )
            stats.append(("Process data", f"{pred.output.get('defect_probability', 0):.0%} defect probability"))
        elif kind == "timeseries":
            sensors = pred.raw_scores.get("sensor_importance", [])[:6]
            if sensors:
                charts["timeseries"] = _hbar_png(
                    [s["sensor"] for s in sensors],
                    [s["importance"] for s in sensors],
                    ["#10b981" if s["importance"] >= 0 else "#ef4444" for s in sensors],
                    "Machine-health sensor attributions (IG)",
                )
            stats.append(("Machine health", f"RUL {pred.raw_scores.get('rul', 0):.0f} · {pred.output.get('risk', 0):.0%} risk"))

    return {
        "inspection": inspection,
        "health_png": _health_png(inspection.health_score, inspection.health_band.value)
        if inspection.health_score is not None
        else None,
        "charts": charts,
        "gradcam_path": gradcam_path,
        "stats": stats,
        "narrative": _narrative(session, inspection),
        "decisions": decisions,
    }


def _build_pdf(data: dict, path: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        Image as RLImage,
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    inspection = data["inspection"]
    styles = getSampleStyleSheet()
    story: list = []

    story.append(Paragraph("SentinelQ Inspection Report", styles["Title"]))
    story.append(
        Paragraph(
            f"Inspection #{inspection.id} · {inspection.category or 'inspection'} · "
            f"status {inspection.status.value} · generated {datetime.now(timezone.utc):%Y-%m-%d %H:%M UTC}",
            styles["Normal"],
        )
    )
    story.append(Spacer(1, 0.4 * cm))

    if data["health_png"]:
        story.append(RLImage(io.BytesIO(data["health_png"]), width=13 * cm, height=3 * cm))
        story.append(Spacer(1, 0.3 * cm))

    if data["stats"]:
        story.append(Paragraph("Model outputs", styles["Heading2"]))
        for name, value in data["stats"]:
            story.append(Paragraph(f"<b>{name}:</b> {value}", styles["Normal"]))
        story.append(Spacer(1, 0.3 * cm))

    if data["gradcam_path"] and Path(data["gradcam_path"]).exists():
        story.append(Paragraph("Vision — Grad-CAM", styles["Heading3"]))
        story.append(RLImage(data["gradcam_path"], width=6 * cm, height=6 * cm))
        story.append(Spacer(1, 0.3 * cm))

    for key, title in (("tabular", "Process data — SHAP"), ("timeseries", "Machine health — Integrated Gradients")):
        if key in data["charts"]:
            story.append(Paragraph(title, styles["Heading3"]))
            story.append(RLImage(io.BytesIO(data["charts"][key]), width=13 * cm, height=6 * cm))
            story.append(Spacer(1, 0.2 * cm))

    narrative = data["narrative"]
    if narrative:
        story.append(Paragraph("Analysis", styles["Heading2"]))
        story.append(Paragraph(f"<b>Root cause:</b> {narrative.get('root_cause', '')}", styles["Normal"]))
        if narrative.get("contributing_factors"):
            story.append(Paragraph("Contributing factors:", styles["Heading4"]))
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(f, styles["Normal"])) for f in narrative["contributing_factors"]],
                    bulletType="bullet",
                )
            )
        if narrative.get("recommendations"):
            story.append(Paragraph("Recommendations:", styles["Heading4"]))
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(r, styles["Normal"])) for r in narrative["recommendations"]],
                    bulletType="bullet",
                )
            )
        if narrative.get("confidence_note"):
            story.append(Spacer(1, 0.2 * cm))
            story.append(Paragraph(f"<i>{narrative['confidence_note']}</i>", styles["Normal"]))
        story.append(Spacer(1, 0.3 * cm))

    if data["decisions"]:
        story.append(Paragraph("Inspector decisions", styles["Heading2"]))
        for decision in data["decisions"]:
            corrected = f" → {decision.corrected_label}" if decision.corrected_label else ""
            note = f" — {decision.note}" if decision.note else ""
            story.append(Paragraph(f"{decision.decision.value}{corrected}{note}", styles["Normal"]))
        story.append(Spacer(1, 0.3 * cm))

    story.append(Spacer(1, 0.5 * cm))
    story.append(
        Paragraph(
            f"Narrative provider: {inspection.llm_provider_used or 'mock'} · "
            "Predictions by SentinelQ deep-learning models (CNN / ANN / LSTM).",
            styles["Italic"],
        )
    )

    SimpleDocTemplate(str(path), pagesize=A4).build(story)


def _build_docx(data: dict, path: Path) -> None:
    from docx import Document
    from docx.shared import Inches

    inspection = data["inspection"]
    document = Document()
    document.add_heading("SentinelQ Inspection Report", level=0)
    document.add_paragraph(
        f"Inspection #{inspection.id} · {inspection.category or 'inspection'} · "
        f"status {inspection.status.value} · generated {datetime.now(timezone.utc):%Y-%m-%d %H:%M UTC}"
    )

    if data["health_png"]:
        document.add_picture(io.BytesIO(data["health_png"]), width=Inches(6.0))

    if data["stats"]:
        document.add_heading("Model outputs", level=1)
        for name, value in data["stats"]:
            document.add_paragraph(f"{name}: {value}", style="List Bullet")

    if data["gradcam_path"] and Path(data["gradcam_path"]).exists():
        document.add_heading("Vision — Grad-CAM", level=2)
        document.add_picture(data["gradcam_path"], width=Inches(2.6))

    for key, title in (("tabular", "Process data — SHAP"), ("timeseries", "Machine health — Integrated Gradients")):
        if key in data["charts"]:
            document.add_heading(title, level=2)
            document.add_picture(io.BytesIO(data["charts"][key]), width=Inches(6.0))

    narrative = data["narrative"]
    if narrative:
        document.add_heading("Analysis", level=1)
        document.add_paragraph(narrative.get("root_cause", ""))
        if narrative.get("contributing_factors"):
            document.add_heading("Contributing factors", level=3)
            for factor in narrative["contributing_factors"]:
                document.add_paragraph(factor, style="List Bullet")
        if narrative.get("recommendations"):
            document.add_heading("Recommendations", level=3)
            for rec in narrative["recommendations"]:
                document.add_paragraph(rec, style="List Bullet")
        if narrative.get("confidence_note"):
            document.add_paragraph(narrative["confidence_note"]).italic = True

    if data["decisions"]:
        document.add_heading("Inspector decisions", level=1)
        for decision in data["decisions"]:
            corrected = f" -> {decision.corrected_label}" if decision.corrected_label else ""
            note = f" — {decision.note}" if decision.note else ""
            document.add_paragraph(f"{decision.decision.value}{corrected}{note}", style="List Bullet")

    document.add_paragraph(
        f"Narrative provider: {inspection.llm_provider_used or 'mock'} · "
        "Predictions by SentinelQ deep-learning models (CNN / ANN / LSTM)."
    ).italic = True

    document.save(str(path))


def generate_report(session: Session, inspection_id: int, fmt: str = "pdf") -> Path:
    if fmt not in {"pdf", "docx"}:
        raise ValueError("format must be 'pdf' or 'docx'")
    data = _gather(session, inspection_id)
    if data is None:
        raise ValueError(f"Inspection {inspection_id} not found.")

    report_dir = get_settings().report_path
    report_dir.mkdir(parents=True, exist_ok=True)
    path = report_dir / f"inspection_{inspection_id}.{fmt}"

    if fmt == "pdf":
        _build_pdf(data, path)
    else:
        _build_docx(data, path)
    return path


def list_reports() -> list[dict]:
    report_dir = get_settings().report_path
    if not report_dir.exists():
        return []
    reports: list[dict] = []
    for path in sorted(report_dir.glob("inspection_*.*")):
        if path.suffix.lstrip(".") not in {"pdf", "docx"}:
            continue
        try:
            inspection_id = int(path.stem.split("_")[1])
        except (IndexError, ValueError):
            continue
        reports.append(
            {
                "inspection_id": inspection_id,
                "format": path.suffix.lstrip("."),
                "size_bytes": path.stat().st_size,
                "download_url": f"/api/v1/reports/{inspection_id}/download?format={path.suffix.lstrip('.')}",
            }
        )
    return reports
